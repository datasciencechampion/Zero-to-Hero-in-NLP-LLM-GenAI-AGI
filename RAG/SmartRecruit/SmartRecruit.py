import streamlit as st
import zipfile
import os
from docx import Document
import PyPDF2
from langchain.vectorstores import Chroma
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.schema import Document
from langchain.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import ConversationalRetrievalChain
from dotenv import load_dotenv
import openai
import json
import re
import smtplib
from email.mime.text import MIMEText


# Load environment variables
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
client = openai.OpenAI(api_key = api_key )

# Set up working directories
WORKING_FOLDER = "working_cvs"
EMBEDDING_FOLDER = "embedding"
CHUNKS_FOLDER = "chunks"

os.makedirs(WORKING_FOLDER, exist_ok=True)
os.makedirs(EMBEDDING_FOLDER, exist_ok=True)
os.makedirs(CHUNKS_FOLDER, exist_ok=True)



# Helper function to read PDF content using PyPDF2
def read_pdf(file):
    try:
        # Use PyPDF2's PdfReader to read the file object directly
        pdf_reader = PyPDF2.PdfReader(file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        return f"Error reading PDF file: {str(e)}"


# Helper function to extract text from Word document
def read_docx(file):
    try:
        doc = Document(file)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return f"Error reading Word file: {str(e)}"

# Response Generator
def generator(user_query,context):
    prompt = f"""
    You are an AI assistant. Below is the content of a CV:
    {context}
    Based on this CV content, answer the following query:
    "{user_query}"
    """
    response = client.chat.completions.create(
        model="gpt-4-turbo",
        messages=[
            {"role": "system",
             "content": "You are a helpful assistant that answers queries based on CV content."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
    )
    #st.subheader("Response:")
    return response.choices[0].message.content

# Summarizer Function
def summarize_cv(context):
    prompt = f"""
    You are an AI assistant. Below is the content of a CV:
    {context}
    Summarize this CV in a few lines capturing the following:
    - Highest education
    - Number of years of experience
    - Current working company
    - Key technical skills
    """
    response = client.chat.completions.create(
        model="gpt-4-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant that summarizes CV content."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.2,
    )
    return response.choices[0].message.content

# Email Sender Function
def extract_email(content):
    email_pattern = r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+'
    match = re.search(email_pattern, content)
    return match.group(0) if match else None

def send_email(to_email, subject, body):
    try:
        smtp_server = os.getenv("SMTP_SERVER")
        smtp_port = os.getenv("SMTP_PORT")
        smtp_user = os.getenv("SMTP_USER")
        smtp_password = os.getenv("SMTP_PASSWORD")

        msg = MIMEText(body)
        msg["Subject"] = subject
        msg["From"] = smtp_user
        msg["To"] = to_email

        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(smtp_user, smtp_password)
            server.sendmail(smtp_user, to_email, msg.as_string())
        return "Email sent successfully!"
    except Exception as e:
        return f"Failed to send email: {str(e)}"

# Function to process CVs and store embeddings
def process_and_embed_cvs():
    documents = []
    for file_name in os.listdir(WORKING_FOLDER):
        file_path = os.path.join(WORKING_FOLDER, file_name)
        if file_name.endswith(".pdf"):
            with open(file_path, "rb") as pdf_file:
                content = read_pdf(pdf_file)  # Pass file object
        elif file_name.endswith(".docx"):
            with open(file_path, "rb") as docx_file:
                content = read_docx(docx_file)
        else:
            continue

        if content.strip():
            # Split the text into chunks
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
            chunks = text_splitter.split_text(content)

            # Create Document objects for embedding and save each chunk with metadata
            for i, chunk in enumerate(chunks):
                # Add to documents for embedding
                document = Document(page_content=chunk, metadata={"file_name": file_name})
                documents.append(document)

                # Write chunk content to a text file
                chunk_file_name = f"{file_name}_chunk_{i + 1}.txt"
                chunk_file_path = os.path.join(CHUNKS_FOLDER, chunk_file_name)
                with open(chunk_file_path, "w", encoding="utf-8") as chunk_file:
                    chunk_file.write(chunk)

                # Write the Document object (content + metadata) to a JSON file
                document_file_name = f"{file_name}_chunk_{i + 1}.json"
                document_file_path = os.path.join(CHUNKS_FOLDER, document_file_name)
                with open(document_file_path, "w", encoding="utf-8") as document_file:
                    json.dump(
                        {"page_content": document.page_content, "metadata": document.metadata},
                        document_file,
                        indent=4,
                        ensure_ascii=False,
                    )

    # Create embeddings and save them to ChromaDB
    embeddings = OpenAIEmbeddings(
        model="text-embedding-3-large",
        dimensions=1024
    )
    db = Chroma(persist_directory=EMBEDDING_FOLDER, embedding_function=embeddings)
    db.add_documents(documents)
    db.persist()

    return "Embeddings successfully created and stored!"


# Streamlit App
def main():
    st.set_page_config(layout="wide")
    st.title("📋 SmartRecruit 💼")

    # Left panel
    with st.sidebar:
        st.header("⚙️ Options 🛠️")

        # Upload multiple files (.pdf and .docx)
        uploaded_files = st.file_uploader(
            "Upload multiple CVs (.pdf or .docx)",
            type=["pdf", "docx"],
            accept_multiple_files=True
        )

        if uploaded_files:
            for uploaded_file in uploaded_files:
                # Save each file to the WORKING_FOLDER
                file_path = os.path.join(WORKING_FOLDER, uploaded_file.name)
                with open(file_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
            st.success(f"{len(uploaded_files)} file(s) successfully uploaded to {WORKING_FOLDER}!")

        # Embed button
        if st.button("🔠➡️🔢 Embed"):
            result = process_and_embed_cvs()
            st.success(result)

    # Main screen
    st.header("🔍 Query CVs 📄")

    embeddings = OpenAIEmbeddings(
        model="text-embedding-3-large",
        dimensions=1024
    )
    db = Chroma(persist_directory=EMBEDDING_FOLDER, embedding_function=embeddings)

    query = st.text_input("Ask a question to filter CVs (e.g., 'Show me all CVs with experience in Python')")
    if query:
        # Perform semantic search
        results = db.similarity_search(query, k=2)  # Top 10 results
        print("********Results*******", results)
        file_names = list(set([result.metadata["file_name"] for result in results]))

        if file_names:
            st.sidebar.header("📂 Filtered CVs 🔍")
            selected_cv = st.sidebar.radio("Select a CV", file_names)

            if selected_cv:
                file_path = os.path.join(WORKING_FOLDER, selected_cv)

                # Read content of the selected CV for display
                if selected_cv.endswith(".pdf"):
                    with open(file_path, "rb") as pdf_file:
                        content = pdf_file.read()  # Read in binary mode for download
                        text_content = read_pdf(pdf_file)  # Extract text for display
                elif selected_cv.endswith(".docx"):
                    with open(file_path, "rb") as docx_file:
                        content = docx_file.read()  # Read in binary mode for download
                        text_content = read_docx(docx_file)  # Extract text for display
                else:
                    text_content = "Unsupported file type."
                    content = None

                # Display CV content and enable follow-up chat
                st.subheader(f"Selected CV: {selected_cv}")
                st.text_area("CV Content", text_content, height=300, disabled=True)

                # Add a download button for the selected CV
                if content:
                    st.download_button(
                        label="Download Selected CV",
                        data=content,
                        file_name=selected_cv,
                        mime="application/octet-stream" if selected_cv.endswith(
                            ".pdf") else "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                    # Summarize button
                    if st.button("Summarize CV"):
                        summary = summarize_cv(text_content)
                        st.subheader("CV Summary:")
                        st.write(summary)



                follow_up_query = st.text_input("Ask a follow-up question about this CV")
                if st.button("Submit Query"):
                    if follow_up_query:
                        response = generator(follow_up_query, text_content)
                        st.subheader("Response:")
                        st.write(response)
                # Send Email button
                email_address = extract_email(text_content)
                if email_address:
                    st.subheader("Send Email")

                    # Pre-filled subject and body
                    email_subject = st.text_input("Email Subject:", value="Your CV is shortlisted for [Position Name]")
                    email_body = st.text_area(
                        "Email Body:",
                        value=f"""Hi,\n\nGreetings from ABC.\n\nWe would like to inform you that your CV is shortlisted for [Position Name].\nWe would like to schedule a first round of interview with you. Please suggest a suitable date and time.\n\nRegards,\n[Your Name/Company Name]"""
                    )

                    if st.button("Send Email"):
                        if email_subject and email_body:
                            email_status = send_email(email_address, email_subject, email_body)
                            st.write(email_status)
                        else:
                            st.warning("Please fill in both the subject and body of the email.")


if __name__ == "__main__":
    main()
